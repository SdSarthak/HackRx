import os
import requests
import tempfile
import time
import asyncio
import hashlib
from datetime import datetime
from typing import List, Dict, Optional, Union, Any
import re

# FastAPI and middleware imports
from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials

# Pydantic models
from pydantic import BaseModel, HttpUrl, Field

# Document processing imports
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# Utilities
import logging
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Additional document loaders (with error handling)
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    logger.warning("python-docx not installed - DOCX support disabled")

try:
    import extract_msg
except ImportError:
    extract_msg = None
    logger.warning("extract-msg not installed - MSG support disabled")

import email
from email.policy import default

# Initialize FastAPI app
app = FastAPI(
    title="HackRX Policy QA API",
    description="Insurance policy Q&A system with explainable AI",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer()

# Environment variables
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
API_KEY = os.getenv("API_KEY", "d229e5eb06d6a264c4cebecd4fb0dc33e6a81c7bfa1f01945f751424fcac1e3a")
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

if not GEMINI_API_KEY:
    raise RuntimeError("GEMINI_API_KEY environment variable not found!")

# Pydantic models
class SourceClause(BaseModel):
    """Model for source clause information."""
    text: str = Field(..., description="Exact text of the clause")
    page_number: Optional[int] = Field(None, description="Page number where clause was found")
    confidence_score: float = Field(..., ge=0.0, le=1.0, description="Relevance confidence score")

class EnhancedAnswer(BaseModel):
    """Enhanced answer model with explainability."""
    answer: str = Field(..., description="The main answer to the question")
    confidence_score: float = Field(..., ge=0.0, le=1.0, description="Overall confidence in the answer")
    source_clauses: List[SourceClause] = Field(default_factory=list, description="Relevant source clauses")
    reasoning: str = Field(..., description="Explanation of how the answer was derived")

class APIRequest(BaseModel):
    """API request model."""
    documents: Union[HttpUrl, List[HttpUrl]] = Field(..., description="Document URL(s) to process")
    questions: List[str] = Field(..., min_items=1, max_items=50, description="Questions to answer")

class APIResponse(BaseModel):
    """API response model."""
    answers: List[str] = Field(..., description="List of answer strings")

# Global variables
vectorstore = None
llm = None
embeddings_model = None

# Authentication and utility functions
def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Token verification."""
    if credentials.credentials != API_KEY:
        logger.warning("Invalid API key attempt")
        raise HTTPException(status_code=401, detail="Invalid API key")
    return credentials.credentials

def validate_file_size(file_size: int) -> bool:
    """Validate file size against limits."""
    return file_size <= MAX_FILE_SIZE

def download_document(url: str) -> tuple[str, str]:
    """Download document from URL and return local file path and type."""
    try:
        response = requests.get(str(url), timeout=30)
        response.raise_for_status()
        
        # Validate file size
        content_length = response.headers.get('content-length')
        if content_length and not validate_file_size(int(content_length)):
            raise HTTPException(status_code=413, detail="File too large")
        
        # Determine file type
        content_type = response.headers.get('content-type', '').lower()
        file_extension = '.pdf'  # default
        
        if 'pdf' in content_type or url.lower().endswith('.pdf'):
            file_extension = '.pdf'
        elif 'docx' in content_type or url.lower().endswith('.docx'):
            file_extension = '.docx'
        elif 'doc' in content_type or url.lower().endswith('.doc'):
            file_extension = '.doc'
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_file.write(response.content)
            return tmp_file.name, file_extension
            
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to download document: {str(e)}")

def load_pdf_document(file_path: str) -> List[Document]:
    """Load PDF document with enhanced metadata."""
    try:
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        
        # Enhance metadata
        for i, doc in enumerate(documents):
            doc.metadata.update({
                'file_type': 'pdf',
                'page_number': i + 1,
                'source_file': os.path.basename(file_path)
            })
        
        return documents
    except Exception as e:
        logger.error(f"Error loading PDF {file_path}: {e}")
        return []

def load_docx_document(file_path: str) -> List[Document]:
    """Load DOCX document with paragraph-level chunking."""
    if DocxDocument is None:
        logger.error("python-docx not installed - cannot process DOCX files")
        return []
    
    try:
        doc = DocxDocument(file_path)
        documents = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                documents.append(Document(
                    page_content=paragraph.text,
                    metadata={
                        'file_type': 'docx',
                        'paragraph_number': i + 1,
                        'source_file': os.path.basename(file_path)
                    }
                ))
        
        return documents
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        return []

def load_document_by_type(file_path: str, file_type: str) -> List[Document]:
    """Load document based on file type."""
    if file_type == '.pdf':
        return load_pdf_document(file_path)
    elif file_type in ['.docx', '.doc']:
        return load_docx_document(file_path)
    else:
        logger.warning(f"Unsupported file type: {file_type}")
        return []

async def initialize_models():
    """Initialize LLM and embeddings models."""
    global llm, embeddings_model
    
    if not embeddings_model:
        embeddings_model = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=GEMINI_API_KEY
        )
    
    if not llm:
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0,
            google_api_key=GEMINI_API_KEY
        )

def create_vectorstore_from_documents(documents: List[Document]) -> FAISS:
    """Create FAISS vectorstore from documents."""
    global embeddings_model
    
    if not embeddings_model:
        embeddings_model = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=GEMINI_API_KEY
        )
    
    return FAISS.from_documents(documents, embeddings_model)

def extract_clause_references(context: str, answer: str) -> List[SourceClause]:
    """Extract specific clause references from context and answer."""
    clauses = []
    
    # Split context into potential clauses
    sentences = re.split(r'[.!?]\s+', context)
    
    for i, sentence in enumerate(sentences):
        if len(sentence.strip()) < 20:  # Skip very short sentences
            continue
            
        # Calculate relevance score based on keyword overlap with answer
        answer_words = set(answer.lower().split())
        sentence_words = set(sentence.lower().split())
        overlap = len(answer_words & sentence_words)
        relevance_score = min(overlap / max(len(answer_words), 1), 1.0)
        
        if relevance_score > 0.2:  # Threshold for relevance
            clauses.append(SourceClause(
                text=sentence.strip(),
                page_number=None,  # Would need to be extracted from metadata
                confidence_score=relevance_score
            ))
    
    # Sort by relevance and return top 3
    clauses.sort(key=lambda x: x.confidence_score, reverse=True)
    return clauses[:3]

def enhance_answer_with_explanations(question: str, raw_answer: str, context: str, 
                                   retrieved_docs: List[Document]) -> EnhancedAnswer:
    """Enhance raw answer with explanations and source tracking."""
    
    # Extract source clauses
    source_clauses = extract_clause_references(context, raw_answer)
    
    # Calculate confidence score based on context relevance
    question_words = set(question.lower().split())
    context_words = set(context.lower().split())
    context_overlap = len(question_words & context_words) / max(len(question_words), 1)
    confidence_score = min(context_overlap * 1.2, 1.0)  # Boost slightly
    
    # Generate reasoning explanation
    reasoning = f"""Answer derived from analysis of {len(retrieved_docs)} document sections. 
Key factors considered: document context relevance ({context_overlap:.2%}), 
source clause alignment, and semantic similarity. Found {len(source_clauses)} relevant clauses."""
    
    return EnhancedAnswer(
        answer=raw_answer,
        confidence_score=confidence_score,
        source_clauses=source_clauses,
        reasoning=reasoning
    )

async def answer_question_enhanced(question: str) -> EnhancedAnswer:
    """Enhanced question answering with explainability."""
    global vectorstore, llm
    
    if not vectorstore or not llm:
        raise HTTPException(status_code=500, detail="Document not processed yet")
    
    try:
        # Retrieve relevant documents
        retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
        docs = retriever.get_relevant_documents(question)
        
        # Create enhanced context
        context_parts = []
        for i, doc in enumerate(docs):
            page_info = f"[Page {doc.metadata.get('page_number', 'N/A')}]" if 'page_number' in doc.metadata else ""
            context_parts.append(f"Source {i+1} {page_info}: {doc.page_content}")
        
        context = "\n\n".join(context_parts)
        
        # Create enhanced prompt
        prompt = f"""You are an expert insurance policy analyst with deep knowledge of insurance terminology, legal clauses, and policy interpretation.

INSTRUCTIONS:
1. Answer the question based ONLY on the provided document context
2. Be precise, specific, and use exact terms from the policy documents
3. Include relevant clause numbers, section references, and specific amounts when available
4. If information is not available in the context, clearly state this
5. Use professional insurance terminology appropriately

CONTEXT FROM POLICY DOCUMENTS:
{context}

QUESTION: {question}

DETAILED ANSWER:"""
        
        # Generate answer
        response = await asyncio.to_thread(llm.invoke, prompt)
        raw_answer = response.content if hasattr(response, 'content') else str(response)
        
        # Enhance answer with explanations
        enhanced_answer = enhance_answer_with_explanations(
            question, raw_answer, context, docs
        )
        
        logger.info(f"Generated enhanced answer for question: {question[:50]}...")
        return enhanced_answer
        
    except Exception as e:
        logger.error(f"Error answering question: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing question: {str(e)}")

# Main API endpoint
@app.post("/hackrx/run", response_model=APIResponse)
async def hackrx_run(request: APIRequest, api_key: str = Depends(verify_token)) -> APIResponse:
    """Main endpoint for HackRX competition - handles document processing and Q&A."""
    start_time = time.time()
    
    try:
        logger.info(f"Processing request with {len(request.questions)} questions")
        
        # Handle multiple document URLs
        document_urls = request.documents if isinstance(request.documents, list) else [request.documents]
        temp_files = []
        all_documents = []
        
        await initialize_models()
        
        # Download and process documents
        for url in document_urls:
            file_path, file_type = download_document(url)
            temp_files.append(file_path)
            
            try:
                # Load document based on type
                documents = load_document_by_type(file_path, file_type)
                if documents:
                    all_documents.extend(documents)
                    logger.info(f"Loaded {len(documents)} sections from {file_type} document")
                else:
                    logger.warning(f"No content found in document: {url}")
                
            except Exception as e:
                logger.error(f"Error processing document {url}: {e}")
                continue
        
        if not all_documents:
            raise HTTPException(status_code=400, detail="No valid documents could be processed")
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300,
            separators=["\n\n", "\n", ". ", ", ", " ", ""]
        )
        splits = text_splitter.split_documents(all_documents)
        
        # Create vector store
        global vectorstore
        vectorstore = create_vectorstore_from_documents(splits)
        logger.info(f"Created vector store with {len(splits)} chunks")
        
        # Process all questions with simple string responses
        answers = []
        
        for i, question in enumerate(request.questions):
            logger.info(f"Processing question {i+1}/{len(request.questions)}: {question[:50]}...")
            
            try:
                enhanced_answer = await answer_question_enhanced(question)
                # Extract just the answer string for the response
                answers.append(enhanced_answer.answer)
                
            except Exception as e:
                logger.error(f"Error processing question {i+1}: {e}")
                # Add error message as string
                answers.append(f"Error processing question: {str(e)}")
        
        processing_time = time.time() - start_time
        
        response = APIResponse(
            answers=answers
        )
        
        logger.info(f"Processing completed in {processing_time:.2f}s")
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in hackrx_run: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")
    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                os.unlink(temp_file)
            except Exception as e:
                logger.warning(f"Failed to cleanup temp file {temp_file}: {e}")

# Health check endpoint (minimal)
@app.get("/")
async def root():
    """Root endpoint."""
    return {
        "message": "HackRX Policy QA API is running",
        "endpoint": "/hackrx/run",
        "status": "ready"
    }

if __name__ == "__main__":
    import uvicorn
    logger.info("Starting HackRX Policy QA API")
    uvicorn.run(app, host="0.0.0.0", port=8000)
